import os
import json
from datetime import datetime
from typing import Dict, List, Union, Optional

import matplotlib
import matplotlib.pyplot as plt
import numpy as np
import networkx as nx
from docx import Document
from docx.shared import Inches

# —— 字体设置，修复中文方块/负号 ——
matplotlib.rcParams['font.sans-serif'] = [
    'PingFang SC','Heiti SC','Hiragino Sans GB','Arial Unicode MS','Noto Sans CJK SC','DejaVu Sans'
]
matplotlib.rcParams['axes.unicode_minus'] = False

ART_DIR = "artifacts"
os.makedirs(ART_DIR, exist_ok=True)

# =======================
# 基础流程图
# =======================

def draw_flow(methods: List[str], out_path: str = os.path.join(ART_DIR, "flow.png")) -> str:
    """绘制流程图：输入→多维判断→方法论匹配→每个方法"""
    G = nx.DiGraph()
    G.add_edge("客户数据输入", "多维度判断")
    G.add_edge("多维度判断", "方法论匹配引擎")
    for m in methods:
        G.add_edge("方法论匹配引擎", m)
    pos = nx.spring_layout(G, seed=42)
    plt.figure(figsize=(8, 6))
    nx.draw(G, pos, with_labels=True, node_size=2200, font_size=10)
    plt.tight_layout()
    plt.savefig(out_path, dpi=150, bbox_inches='tight')
    plt.close()
    return out_path

# =======================
# BCG 四象限
# =======================

def _normalize_bcg_points(points: Union[List[Dict], 'np.ndarray']):
    """将输入标准化为列表[{name,growth,share}]，容忍字段别名。"""
    norm = []
    if points is None:
        return norm
    if isinstance(points, list):
        for p in points:
            if not isinstance(p, dict):
                continue
            name = p.get('BU') or p.get('name') or p.get('业务') or p.get('业务单元') or 'BU'
            growth = p.get('market_growth') or p.get('growth') or p.get('增长率')
            share = p.get('rel_share') or p.get('share') or p.get('相对份额')
            try:
                growth = float(growth) if growth is not None else None
                share = float(share) if share is not None else None
            except Exception:
                growth, share = None, None
            if growth is not None and share is not None:
                norm.append({'name': name, 'growth': growth, 'share': share})
    return norm


def draw_bcg(points: Optional[List[Dict]] = None,
             out_path: str = os.path.join(ART_DIR, "bcg.png")) -> Optional[str]:
    """绘制 BCG 四象限图。
    points: [{name, growth(市场增长率%), share(相对份额)}]
    """
    pts = _normalize_bcg_points(points)
    if not pts:
        return None

    # 阈值：增长率=10% 作为高/低，份额=1.0 作为高/低（可按需调整）
    g_thr, s_thr = 10.0, 1.0

    plt.figure(figsize=(7, 6))
    ax = plt.gca()
    ax.axvline(s_thr, linestyle='--')
    ax.axhline(g_thr, linestyle='--')
    ax.set_xlabel('相对市场份额')
    ax.set_ylabel('市场增长率 (%)')
    ax.set_title('BCG 矩阵')

    for p in pts:
        x, y = p['share'], p['growth']
        ax.scatter(x, y, s=120)
        ax.text(x + 0.03, y + 0.5, p['name'])

    # 象限标签
    ax.text(s_thr + 0.1, g_thr + 0.1, '明星', fontsize=11)
    ax.text(0.1, g_thr + 0.1, '问题', fontsize=11)
    ax.text(s_thr + 0.1, 0.1, '现金牛', fontsize=11)
    ax.text(0.1, 0.1, '瘦狗', fontsize=11)

    plt.tight_layout()
    plt.savefig(out_path, dpi=150, bbox_inches='tight')
    plt.close()
    return out_path

# =======================
# 波特五力 雷达图
# =======================

def _normalize_porter_scores(scores: Union[Dict, List[Dict]]):
    """标准化为 dict[label->score]，得分范围建议 1–5。"""
    if scores is None:
        return {}
    if isinstance(scores, dict):
        return scores
    if isinstance(scores, list):
        result = {}
        for row in scores:
            if isinstance(row, dict):
                k = row.get('力量') or row.get('factor') or row.get('name')
                v = row.get('评分') or row.get('score') or row.get('value')
                if k is not None and v is not None:
                    try:
                        result[str(k)] = float(v)
                    except Exception:
                        pass
        return result
    return {}


def draw_porter_radar(scores: Union[Dict, List[Dict], None] = None,
                      out_path: str = os.path.join(ART_DIR, "porter_radar.png")) -> Optional[str]:
    data = _normalize_porter_scores(scores)
    if not data:
        return None
    labels = list(data.keys())
    values = [data[k] for k in labels]
    # 闭合雷达
    labels.append(labels[0])
    values.append(values[0])

    angles = np.linspace(0, 2 * np.pi, len(labels), endpoint=False)
    fig = plt.figure(figsize=(6, 6))
    ax = plt.subplot(111, polar=True)
    ax.plot(angles, values)
    ax.fill(angles, values, alpha=0.2)
    ax.set_thetagrids(angles * 180/np.pi, labels)
    ax.set_title('波特五力雷达图 (1-5)')
    plt.tight_layout()
    plt.savefig(out_path, dpi=150, bbox_inches='tight')
    plt.close()
    return out_path

# =======================
# 轻量数据加载（若 data/ 目录有示例文件）
# =======================

def _load_optional_samples():
    bcg_points = None
    porter_scores = None
    # 支持多种文件名
    for fn in ["data/bcg.json", "data/bcg_sample.json", "data/byd_bcg.json"]:
        if os.path.exists(fn):
            try:
                with open(fn, 'r', encoding='utf-8') as f:
                    bcg_points = json.load(f)
                break
            except Exception:
                pass
    for fn in ["data/porter.json", "data/porter_sample.json", "data/byd_porter.json"]:
        if os.path.exists(fn):
            try:
                with open(fn, 'r', encoding='utf-8') as f:
                    porter_scores = json.load(f)
                break
            except Exception:
                pass
    return bcg_points, porter_scores

# =======================
# 报告生成（Markdown / Word）
# =======================

def _features_md(feats: Dict) -> str:
    lines = []
    for k, v in feats.items():
        if isinstance(v, dict):
            lines.append(f"- **{k}**:")
            for kk, vv in v.items():
                lines.append(f"  - {kk}: {vv}")
        else:
            lines.append(f"- **{k}**: {v}")
    return "\n".join(lines)


def save_markdown_report(methods: List[str],
                         feats: Dict,
                         triggers: List[str],
                         models_data: Dict,
                         out_path: str = os.path.join(ART_DIR, "report.md"),
                         bcg_points: Optional[List[Dict]] = None,
                         porter_scores: Optional[Union[Dict, List[Dict]]] = None) -> str:
    """导出 Markdown 报告（含方法简介、触发依据、可用时插入 BCG/五力图）。"""
    # 兜底加载示例
    if bcg_points is None and porter_scores is None:
        bcg_points, porter_scores = _load_optional_samples()

    lines = []
    lines.append(f"# 战略分析报告\n\n生成时间：{datetime.now().isoformat(timespec='seconds')}\n")
    lines.append("## 推荐方法论\n" + ("、".join(methods) if methods else "（尚未运行分析）"))
    if triggers:
        lines.append("\n## 触发依据\n" + "\n".join([f"- {t}" for t in triggers]))
    lines.append("\n## 关键特征\n" + _features_md(feats))

    # 方法论简介
    if methods and models_data:
        lines.append("\n## 方法论说明")
        for m in methods:
            meta = models_data.get(m) or models_data.get(m.replace('Porter Five Forces', '波特五力')) or {}
            lines.append(f"\n### {m}")
            if meta.get('简介'):
                lines.append(f"**简介**：{meta['简介']}")
            if meta.get('应用场景'):
                lines.append(f"**应用场景**：{meta['应用场景']}")
            steps = meta.get('分析步骤') or []
            if steps:
                lines.append("**分析步骤**：")
                for s in steps:
                    lines.append(f"- {s}")

    # 流程图
    flow_path = os.path.join(ART_DIR, "flow.png")
    if os.path.exists(flow_path):
        lines.append(f"\n## 分析流程图\n![流程图]({flow_path})\n")

    # BCG & 五力
    bcg_path = draw_bcg(bcg_points) if bcg_points else None
    if bcg_path:
        lines.append(f"\n## BCG 矩阵\n![BCG]({bcg_path})\n")
    porter_path = draw_porter_radar(porter_scores) if porter_scores else None
    if porter_path:
        lines.append(f"\n## 波特五力雷达图\n![Porter]({porter_path})\n")

    with open(out_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(lines))
    return out_path


def save_docx_report(methods: List[str],
                     feats: Dict,
                     triggers: List[str],
                     models_data: Dict,
                     out_path: str = os.path.join(ART_DIR, "report.docx"),
                     bcg_points: Optional[List[Dict]] = None,
                     porter_scores: Optional[Union[Dict, List[Dict]]] = None) -> str:
    """导出 Word 报告（含流程图 + BCG + 五力雷达图（如有数据））。"""
    # 兜底加载示例
    if bcg_points is None and porter_scores is None:
        bcg_points, porter_scores = _load_optional_samples()

    doc = Document()
    doc.add_heading('战略分析报告', 0)
    doc.add_paragraph(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))

    doc.add_heading('一、推荐方法论', level=1)
    doc.add_paragraph("、".join(methods) if methods else "（尚未运行分析）")

    if triggers:
        doc.add_heading('二、触发依据', level=1)
        for t in triggers:
            doc.add_paragraph(t, style='List Bullet')

    doc.add_heading('三、关键特征', level=1)
    for k, v in feats.items():
        if isinstance(v, dict):
            doc.add_paragraph(f"{k}")
            for kk, vv in v.items():
                doc.add_paragraph(f"{kk}: {vv}", style='List Bullet')
        else:
            doc.add_paragraph(f"{k}: {v}")

    # 方法论说明
    if methods and models_data:
        doc.add_heading('四、方法论说明', level=1)
        for m in methods:
            meta = models_data.get(m) or models_data.get(m.replace('Porter Five Forces', '波特五力')) or {}
            doc.add_heading(m, level=2)
            if meta.get('简介'):
                doc.add_paragraph(f"简介：{meta['简介']}")
            if meta.get('应用场景'):
                doc.add_paragraph(f"应用场景：{meta['应用场景']}")
            steps = meta.get('分析步骤') or []
            if steps:
                doc.add_paragraph('分析步骤：')
                for s in steps:
                    doc.add_paragraph(s, style='List Bullet')

    # 图片：流程图 / BCG / 五力
    flow_path = draw_flow(methods or [])
    try:
        doc.add_heading('五、分析流程图', level=1)
        doc.add_picture(flow_path, width=Inches(6))
    except Exception:
        pass

    bcg_path = draw_bcg(bcg_points)
    if bcg_path:
        try:
            doc.add_heading('六、BCG 矩阵', level=1)
            doc.add_picture(bcg_path, width=Inches(6))
        except Exception:
            pass

    porter_path = draw_porter_radar(porter_scores)
    if porter_path:
        try:
            doc.add_heading('七、波特五力雷达图', level=1)
            doc.add_picture(porter_path, width=Inches(6))
        except Exception:
            pass

    doc.save(out_path)
    return out_path
